"""Build the article revision with every change from the original in bold."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "TABD-PC-FINAL_REVISADO_MAYORES.docx"
OUTPUT = ROOT / "TABD-PC-FINAL_REVISADO_MAYORES_GEMINI_NEGRITA.docx"


REPLACEMENTS = {
    "Abstract.": (
        "Abstract. El aprendizaje autodirigido de programación competitiva presenta dificultades para "
        "estudiantes principiantes, debido a la dispersión de recursos y la necesidad de relacionar "
        "restricciones, técnicas algorítmicas y editoriales. Este artículo presenta el desarrollo y una "
        "evaluación preliminar de un sistema basado en RAG, PageIndex y Hybrid Tree Search. La evaluación "
        "principal comparó Page Nodes y PageIndex Chunks mediante 40 consultas sobre 40 de los 80 problemas de "
        "Codeforces del corpus. Page Nodes obtuvo 0.86 en Faithfulness y 0.86 en Answer Relevancy; PageIndex "
        "Chunks alcanzó 0.83 y 0.86. Para reducir la dependencia del juez original, se realizó una auditoría "
        "cruzada sobre 16 respuestas persistidas con Gemini 2.5 Flash, temperatura 0 y evidencia oficial "
        "recuperada. Gemini obtuvo 0.71 en Faithfulness y 0.94 en Answer Relevancy, con 56.25 % y 81.25 % de "
        "casos sobre sus respectivos umbrales. La diferencia entre jueces muestra que la estimación de fidelidad "
        "es sensible al modelo evaluador. Los resultados se interpretan como evidencia técnica preliminar; el "
        "protocolo de revisión humana fue preparado, pero aún no ha sido aplicado."
    ),
    "El artículo se organiza en siete secciones:": (
        "El artículo se organiza en siete secciones: introducción; conceptos y definiciones; revisión de la "
        "literatura; conceptos de desarrollo; metodología y arquitectura; resultados y discusión; y conclusiones "
        "con trabajo futuro."
    ),
    "La investigación sigue un enfoque aplicativo": (
        "La investigación sigue un enfoque aplicativo con evaluación experimental exploratoria. Se implementa "
        "un prototipo RAG y se comparan dos representaciones de recuperación mediante 40 consultas. Como "
        "comprobación adicional, 16 respuestas persistidas se someten a una auditoría cruzada con un juez de un "
        "proveedor distinto. El estudio no incorpora grupo de control ni participantes humanos; por ello, el "
        "análisis se limita al comportamiento técnico del prototipo y no permite afirmar todavía una mejora del "
        "aprendizaje."
    ),
    "Sprint 1:": (
        "Sprint 1: Recolección y procesamiento del corpus de programación competitiva. Se empleó un script de "
        "Python para obtener metadatos y recuperar editoriales oficiales mediante solicitudes controladas y "
        "cache. El módulo admite Codeforces y AtCoder; sin embargo, el corpus evaluado está formado por 80 "
        "problemas de Codeforces correspondientes a concursos Division 4 recientes. El contenido se conservó "
        "junto con su URL y estado de extracción, sin completar artificialmente las editoriales no recuperadas."
    ),
    "Sprint 3:": (
        "Sprint 3: Implementación del baseline de recuperación vectorial. Se utilizó Page Nodes como baseline "
        "plano. El texto de cada nodo se representó con TF-IDF de unigramas y bigramas, reducción TruncatedSVD "
        "de hasta 128 dimensiones y normalización L2; la recuperación se calculó mediante producto interno, "
        "equivalente a similitud coseno para vectores normalizados. Los mismos vectores fueron consultados en "
        "matriz local, FAISS y ChromaDB, separando el efecto del backend del efecto de la organización jerárquica."
    ),
    "Sprint 5:": (
        "Sprint 5: Construcción del conjunto de evaluación principal. Se seleccionaron 40 problemas del corpus "
        "de 80 mediante un muestreo determinista distribuido por dificultad. Se formularon diez consultas para "
        "cada categoría: observación y algoritmo, restricciones y complejidad, prueba de corrección y riesgos de "
        "implementación. Page Nodes y PageIndex Chunks recibieron las mismas consultas, con recuperación global "
        "top-5 y sin restringir la búsqueda al problema esperado. GPT-4o-mini generó las respuestas y produjo los "
        "puntajes de la evaluación principal."
    ),
    "Ambos indicadores se calcularon": (
        "En la evaluación principal, GPT-4o-mini, con temperatura 0, generó las respuestas y calculó Faithfulness "
        "y Answer Relevancy. Para controlar el riesgo de auto-preferencia, se añadió una auditoría independiente "
        "con Gemini 2.5 Flash, también con temperatura 0. Gemini recibió únicamente la consulta, la respuesta "
        "persistida y el contexto oficial recuperado. Los puntajes de ambos jueces se reportan por separado y no "
        "se promedian, porque representan criterios de evaluación distintos."
    ),
    "El desempeño del sistema se evaluó": (
        "El desempeño principal se evaluó con 40 consultas sobre 40 problemas, equivalentes al 50 % del corpus. "
        "Se compararon Page Nodes y PageIndex Chunks bajo recuperación global top-5, usando los umbrales 0.80 "
        "para Faithfulness y 0.85 para Answer Relevancy. La Tabla 1 presenta los resultados del juez original. "
        "Posteriormente, Gemini 2.5 Flash auditó las 16 respuestas que permanecían persistidas, ocho por modo de "
        "indexación; la Tabla 2 presenta esta validación cruzada independiente."
    ),
    "Table 1.": "Table 1. Evaluación principal entre Page Nodes y PageIndex Chunks (n = 40)",
    "Page Nodes alcanzó": (
        "En la evaluación principal, Page Nodes alcanzó 0.86 en Faithfulness, desviación estándar de 0.086 y "
        "90 % de casos sobre el umbral. PageIndex Chunks obtuvo 0.83, desviación estándar de 0.126 y 88 %. En "
        "Answer Relevancy, ambos modos promediaron 0.86; Page Nodes presentó una desviación estándar de 0.107 y "
        "82 % de cumplimiento, mientras PageIndex Chunks obtuvo 0.088 y 78 %. En esta muestra, la representación "
        "jerárquica no superó al baseline plano."
    ),
    "Para completar la comparación": (
        "Para completar la comparación con el baseline del Sprint 3, una prueba exploratoria de recuperación "
        "sobre cinco consultas produjo el mismo ranking con matriz local, FAISS y ChromaDB: Precision@8 = 0.05, "
        "Recall@8 = 0.04 y MRR = 0.1667. Las latencias medias fueron 10.277 ms, 74.251 ms y 1729.670 ms. La "
        "coincidencia se explica porque los tres backends recibieron los mismos vectores TF-IDF+SVD normalizados "
        "y ejecutaron una búsqueda exacta equivalente a coseno. Esta prueba compara infraestructura y latencia, "
        "pero no demuestra una mejora de calidad atribuible a Hybrid Tree Search."
    ),
    "En conjunto, los resultados": (
        "La auditoría con Gemini produjo un promedio global de 0.7063 en Faithfulness, desviación estándar de "
        "0.2893 y 56.25 % de respuestas sobre 0.80. Para Answer Relevancy obtuvo 0.9406, desviación estándar de "
        "0.0852 y 81.25 % sobre 0.85. La diferencia absoluta media entre jueces fue 0.2469 para Faithfulness y "
        "0.0906 para Answer Relevancy; el acuerdo al clasificar casos sobre o bajo el umbral fue 56.25 % y "
        "81.25 %. Gemini fue más estricto al verificar el sustento factual, lo que confirma que los resultados "
        "dependen del modelo evaluador y evita presentar un único juez como referencia definitiva."
    ),
    "La evaluación automática mediante LLM-as-a-Judge": (
        "La evaluación principal sobre 40 problemas mostró resultados cercanos entre las representaciones, pero "
        "la auditoría independiente con Gemini reveló una menor fidelidad en la submuestra persistida. Esta "
        "triangulación no demuestra superioridad de PageIndex ni eficacia educativa; sí aporta evidencia de que "
        "el prototipo genera respuestas generalmente relevantes y de que la fidelidad debe revisarse con más de "
        "un juez. La contribución validada en esta etapa es la construcción del corpus, su estructuración y un "
        "flujo experimental reproducible con evaluación cruzada."
    ),
    "Como trabajo futuro se plantea": (
        "Como trabajo futuro se plantea completar la validación humana preparada en ocho casos, con dos "
        "revisores independientes con experiencia en programación competitiva y escalas de 1 a 5 para fidelidad, "
        "relevancia y utilidad pedagógica. El acuerdo se calculará mediante Cohen Kappa ponderado y se comparará "
        "con Gemini mediante correlación de Spearman. También se evaluará el recorrido completo de Hybrid Tree "
        "Search frente al baseline vectorial con el mismo ground truth y se ampliará la auditoría independiente a "
        "las 40 respuestas cuando todas permanezcan almacenadas."
    ),
}


SPRINT6_TEXT = (
    "Sprint 6: Auditoría cruzada y protocolo de validación externa. Se reutilizaron 16 respuestas persistidas, "
    "ocho de Page Nodes y ocho de PageIndex Chunks, sin volver a generarlas. Gemini 2.5 Flash actuó como juez "
    "independiente y evaluó cada respuesta contra el contexto oficial recuperado. Se calcularon diferencia "
    "absoluta media, correlación y acuerdo de clasificación respecto del juez original. Además, se preparó un "
    "protocolo ciego de ocho casos para dos revisores humanos; sus calificaciones permanecen pendientes y no se "
    "presentan como resultados ejecutados."
)

O5_TEXT = (
    "La validez externa se fortaleció sin construir respuestas ideales sintéticas: las 16 respuestas fueron "
    "contrastadas con enunciados y editoriales oficiales, y todos los contextos recuperados correspondieron al "
    "problema consultado, con una precisión por identificador de 1.00. La combinación de evidencia oficial y un "
    "juez de otro proveedor reduce la dependencia del autor y del modelo generador. No obstante, esta es una "
    "validación externa automatizada; el protocolo humano ya está definido, pero hasta completar sus valoraciones "
    "no se afirma utilidad educativa observada en usuarios reales."
)

TABLE2_DATA = [
    ["Modo / Métrica", "n", "Promedio", "Desv. estándar", "% sobre umbral"],
    ["Page Nodes\nFaithfulness", "8", "0.64", "0.356", "50.0 %"],
    ["Page Nodes\nAnswer Relevancy", "8", "0.92", "0.093", "75.0 %"],
    ["PageIndex Chunks\nFaithfulness", "8", "0.78", "0.177", "62.5 %"],
    ["PageIndex Chunks\nAnswer Relevancy", "8", "0.96", "0.070", "87.5 %"],
]


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_with_bold(paragraph: Paragraph, text: str) -> None:
    body_run = next((run for run in paragraph.runs if run.text), None)
    rpr = deepcopy(body_run._element.rPr) if body_run is not None and body_run._element.rPr is not None else None
    clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    if rpr is not None:
        if run._element.rPr is not None:
            run._element.remove(run._element.rPr)
        run._element.insert(0, rpr)
    run.bold = True


def insert_after(anchor: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    if anchor._p.pPr is not None:
        new_p.append(deepcopy(anchor._p.pPr))
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    run = paragraph.add_run(text)
    run.bold = True
    return paragraph


def set_cell_margins(cell, value: int = 55) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_data(table, rows: list[list[str]], widths: list[float]) -> None:
    width_values = [Inches(value) for value in widths]
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, width_values):
        grid_col.w = width
    while len(table.rows) < len(rows):
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, (row, values) in enumerate(zip(table.rows, rows)):
        for col_index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = width_values[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            replace_with_bold(paragraph, value)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(7.2 if row_index else 7.4)
                run.bold = True
            for extra in cell.paragraphs[1:]:
                extra._element.getparent().remove(extra._element)


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"No se encontró el párrafo: {prefix}")


def main() -> None:
    document = Document(SOURCE)
    for prefix, replacement in REPLACEMENTS.items():
        replace_with_bold(find_paragraph(document, prefix), replacement)

    sprint5 = find_paragraph(document, "Sprint 5:")
    insert_after(sprint5, SPRINT6_TEXT)

    limitations = find_paragraph(document, "La auditoría con Gemini")
    insert_after(limitations, O5_TEXT)

    primary_table = document.tables[0]
    for row in primary_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    table2 = document.add_table(rows=5, cols=5)
    set_table_data(table2, TABLE2_DATA, [1.30, 0.38, 0.70, 1.05, 0.95])
    caption = document.add_paragraph()
    replace_with_bold(caption, "Table 2. Auditoría independiente con Gemini 2.5 Flash (n = 16 respuestas)")
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

    primary_table._tbl.addnext(caption._p)
    caption._p.addnext(table2._tbl)

    document.core_properties.subject = "Revisión O1-O5 con auditoría Gemini; cambios en negrita"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
