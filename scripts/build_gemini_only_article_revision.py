"""Build the article revision using Gemini as the only automatic judge.

Every paragraph and table cell changed relative to the reviewed source document
is formatted in bold so the author can identify the revision immediately.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "TABD-PC-FINAL_REVISADO_MAYORES.docx"
OUTPUT = ROOT / "TABD-PC-FINAL_REVISADO_GEMINI_N16_NEGRITA.docx"


REPLACEMENTS = {
    "Abstract.": (
        "Abstract. El aprendizaje autodirigido de programación competitiva presenta dificultades para "
        "estudiantes principiantes, debido a la dispersión de recursos y la necesidad de relacionar "
        "restricciones, técnicas algorítmicas y editoriales. Este artículo presenta el desarrollo y una "
        "evaluación preliminar de un sistema basado en RAG, PageIndex y Hybrid Tree Search. Se evaluaron ocho "
        "consultas aplicadas a dos representaciones de recuperación, Page Nodes y PageIndex Chunks, con un total "
        "de 16 respuestas persistidas. Gemini 2.5 Flash actuó como único juez automático, con temperatura 0 y "
        "acceso a la consulta, la respuesta y el contexto oficial recuperado. El promedio global fue 0.71 en "
        "Faithfulness y 0.94 en Answer Relevancy; el 56.25 % y el 81.25 % de las respuestas superaron los umbrales "
        "de 0.80 y 0.85, respectivamente. Los resultados se interpretan como evidencia técnica preliminar; el "
        "protocolo de revisión humana fue preparado, pero aún no ha sido aplicado."
    ),
    "El artículo se organiza en siete secciones:": (
        "El artículo se organiza en siete secciones: introducción; conceptos y definiciones; revisión de la "
        "literatura; conceptos de desarrollo; metodología y arquitectura; resultados y discusión; y conclusiones "
        "con trabajo futuro."
    ),
    "La investigación sigue un enfoque aplicativo": (
        "La investigación sigue un enfoque aplicativo con evaluación experimental exploratoria. Se implementa "
        "un prototipo RAG y se comparan dos representaciones de recuperación mediante ocho consultas aplicadas a "
        "cada representación, para un total de 16 respuestas. Gemini 2.5 Flash se utiliza como único evaluador "
        "automático. El estudio no incorpora grupo de control ni participantes humanos; por ello, el análisis se "
        "limita al comportamiento técnico del prototipo y no permite afirmar todavía una mejora del aprendizaje."
    ),
    "Sprint 1:": (
        "Sprint 1: Recolección y procesamiento del corpus de programación competitiva. Se empleó un script de "
        "Python para obtener metadatos y recuperar editoriales oficiales mediante solicitudes controladas y "
        "cache. El módulo admite Codeforces y AtCoder; sin embargo, el corpus evaluado está formado por 80 "
        "problemas de Codeforces correspondientes a concursos Division 4 recientes. El contenido se conservó "
        "junto con su URL y estado de extracción, sin completar artificialmente las editoriales no recuperadas."
    ),
    "Sprint 3:": (
        "Sprint 3: Implementación del baseline de recuperación vectorial. Se utilizó Page Nodes como baseline "
        "plano. El texto de cada nodo se representó con TF-IDF de unigramas y bigramas, reducción TruncatedSVD "
        "de hasta 128 dimensiones y normalización L2; la recuperación se calculó mediante producto interno, "
        "equivalente a similitud coseno para vectores normalizados. Los mismos vectores fueron consultados en "
        "matriz local, FAISS y ChromaDB, separando el efecto del backend del efecto de la organización jerárquica."
    ),
    "Sprint 5:": (
        "Sprint 5: Construcción del conjunto de evaluación automática. Se formularon ocho consultas sobre "
        "problemas del corpus y se aplicaron a Page Nodes y PageIndex Chunks bajo las mismas condiciones de "
        "recuperación top-5. GPT-4o-mini generó las respuestas y estas se almacenaron junto con los identificadores "
        "de los nodos recuperados. El diseño produjo 16 respuestas persistidas, ocho por representación, que "
        "constituyen la muestra evaluada en este artículo."
    ),
    "Ambos indicadores se calcularon": (
        "Faithfulness y Answer Relevancy se calcularon mediante Gemini 2.5 Flash como único juez automático, con "
        "temperatura 0. Para cada caso, el evaluador recibió la consulta, la respuesta persistida y el contexto "
        "oficial recuperado. Faithfulness mide el grado en que las afirmaciones de la respuesta están respaldadas "
        "por el contexto; Answer Relevancy estima si la respuesta atiende directamente la consulta planteada."
    ),
    "El desempeño del sistema se evaluó": (
        "El desempeño se evaluó sobre 16 respuestas persistidas: ocho obtenidas mediante Page Nodes y ocho "
        "mediante PageIndex Chunks. Ambas representaciones utilizaron recuperación top-5. Se establecieron los "
        "umbrales de 0.80 para Faithfulness y 0.85 para Answer Relevancy. La Tabla 1 presenta exclusivamente los "
        "resultados calculados por Gemini 2.5 Flash."
    ),
    "Table 1.": "Table 1. Evaluación automática con Gemini 2.5 Flash (n = 16 respuestas)",
    "Page Nodes alcanzó": (
        "Gemini asignó a Page Nodes un promedio de 0.64 en Faithfulness, desviación estándar de 0.356 y 50.0 % "
        "de casos sobre el umbral. PageIndex Chunks obtuvo 0.78, desviación estándar de 0.177 y 62.5 %. En Answer "
        "Relevancy, Page Nodes alcanzó 0.92, desviación estándar de 0.093 y 75.0 % de cumplimiento; PageIndex "
        "Chunks obtuvo 0.96, desviación estándar de 0.070 y 87.5 %. En esta muestra, PageIndex Chunks presentó "
        "valores superiores en ambas métricas, aunque el tamaño reducido exige interpretar la diferencia como "
        "evidencia preliminar."
    ),
    "Para completar la comparación": (
        "Para completar la comparación con el baseline del Sprint 3, una prueba exploratoria de recuperación "
        "sobre cinco consultas produjo el mismo ranking con matriz local, FAISS y ChromaDB: Precision@8 = 0.05, "
        "Recall@8 = 0.04 y MRR = 0.1667. Las latencias medias fueron 10.277 ms, 74.251 ms y 1729.670 ms. La "
        "coincidencia se explica porque los tres backends recibieron los mismos vectores TF-IDF+SVD normalizados "
        "y ejecutaron una búsqueda exacta equivalente a coseno. Esta prueba compara infraestructura y latencia, "
        "pero no demuestra una mejora de calidad atribuible a Hybrid Tree Search."
    ),
    "En conjunto, los resultados": (
        "En conjunto, Gemini produjo un promedio global de 0.7063 en Faithfulness, desviación estándar de 0.2893 "
        "y 56.25 % de respuestas sobre el umbral de 0.80. Para Answer Relevancy obtuvo 0.9406, desviación estándar "
        "de 0.0852 y 81.25 % sobre el umbral de 0.85. Los resultados indican que las respuestas fueron, en general, "
        "pertinentes para las consultas, mientras que el sustento factual mostró mayor variación entre casos."
    ),
    "La evaluación automática mediante LLM-as-a-Judge": (
        "La evaluación automática con Gemini sobre 16 respuestas mostró una relevancia alta y una fidelidad más "
        "variable. PageIndex Chunks obtuvo mayores promedios que Page Nodes en esta muestra, pero el número de "
        "casos no permite establecer superioridad estadística ni eficacia educativa. La contribución validada en "
        "esta etapa es la construcción del corpus, su estructuración y un flujo experimental reproducible que "
        "conserva consulta, contexto recuperado, respuesta y puntuación."
    ),
    "Como trabajo futuro se plantea": (
        "Como trabajo futuro se plantea completar la validación humana preparada en ocho casos, con dos "
        "revisores independientes con experiencia en programación competitiva y escalas de 1 a 5 para fidelidad, "
        "relevancia y utilidad pedagógica. El acuerdo se calculará mediante Cohen Kappa ponderado y las "
        "valoraciones humanas se contrastarán con los puntajes de Gemini mediante correlación de Spearman. "
        "También se evaluará el recorrido completo de Hybrid Tree Search frente al baseline vectorial con el "
        "mismo ground truth y se ampliará progresivamente la muestra automática con respuestas almacenadas."
    ),
}


SPRINT6_TEXT = (
    "Sprint 6: Evaluación automática y protocolo de validación externa. Se reutilizaron 16 respuestas "
    "persistidas, ocho de Page Nodes y ocho de PageIndex Chunks, sin volver a generarlas. Gemini 2.5 Flash actuó "
    "como único juez automático y evaluó cada respuesta contra el contexto oficial recuperado. Además, se "
    "preparó un protocolo ciego de ocho casos para dos revisores humanos; sus calificaciones permanecen "
    "pendientes y no se presentan como resultados ejecutados."
)

O5_TEXT = (
    "La validez externa se fortaleció mediante el uso de enunciados y editoriales oficiales como evidencia de "
    "referencia. En los 16 casos, los contextos recuperados correspondieron al problema consultado, con una "
    "precisión por identificador de 1.00. Gemini evaluó las respuestas sin intervenir en su generación. No "
    "obstante, esta sigue siendo una validación automatizada; hasta completar el protocolo humano no se afirma "
    "utilidad educativa observada en usuarios reales."
)

TABLE1_DATA = [
    ["Modo / Métrica", "n", "Promedio", "Desv. estándar", "% sobre umbral"],
    ["Page Nodes\nFaithfulness", "8", "0.64", "0.356", "50.0 %"],
    ["Page Nodes\nAnswer Relevancy", "8", "0.92", "0.093", "75.0 %"],
    ["PageIndex Chunks\nFaithfulness", "8", "0.78", "0.177", "62.5 %"],
    ["PageIndex Chunks\nAnswer Relevancy", "8", "0.96", "0.070", "87.5 %"],
]


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_with_bold(paragraph: Paragraph, text: str) -> None:
    body_run = next((run for run in paragraph.runs if run.text), None)
    rpr = deepcopy(body_run._element.rPr) if body_run is not None and body_run._element.rPr is not None else None
    clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    if rpr is not None:
        if run._element.rPr is not None:
            run._element.remove(run._element.rPr)
        run._element.insert(0, rpr)
    run.bold = True


def insert_after(anchor: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    if anchor._p.pPr is not None:
        new_p.append(deepcopy(anchor._p.pPr))
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    paragraph.add_run(text).bold = True
    return paragraph


def set_cell_margins(cell, value: int = 55) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_data(table, rows: list[list[str]], widths: list[float]) -> None:
    width_values = [Inches(value) for value in widths]
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, width_values):
        grid_col.w = width
    while len(table.rows) < len(rows):
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, (row, values) in enumerate(zip(table.rows, rows)):
        for col_index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = width_values[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            replace_with_bold(paragraph, value)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(7.2 if row_index else 7.4)
                run.bold = True
            for extra in cell.paragraphs[1:]:
                extra._element.getparent().remove(extra._element)


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"No se encontró el párrafo: {prefix}")


def main() -> None:
    document = Document(SOURCE)
    for prefix, replacement in REPLACEMENTS.items():
        replace_with_bold(find_paragraph(document, prefix), replacement)

    sprint5 = find_paragraph(document, "Sprint 5:")
    insert_after(sprint5, SPRINT6_TEXT)

    results = find_paragraph(document, "En conjunto, Gemini")
    insert_after(results, O5_TEXT)

    set_table_data(document.tables[0], TABLE1_DATA, [1.30, 0.38, 0.70, 1.05, 0.95])

    document.core_properties.subject = "Revisión O1-O5; evaluación Gemini n=16; cambios en negrita"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
